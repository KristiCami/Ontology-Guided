import json
import logging
import pytest

from ontology_guided.data_loader import DataLoader


def test_demo_txt_loading_and_preprocessing():
    loader = DataLoader()
    texts = list(loader.load_requirements(["demo.txt"]))
    assert len(texts) == 1
    sentences = []
    for t in texts:
        sentences.extend(loader.preprocess_text(t))
    assert sentences == [
        "The ATM must log all user transactions after card insertion.",
    ]


def test_docx_loading_and_preprocessing(tmp_path):
    from docx import Document

    doc = Document()
    doc.add_paragraph("First requirement.")
    doc.add_paragraph("Second requirement.")
    file_path = tmp_path / "reqs.docx"
    doc.save(file_path)

    loader = DataLoader()
    lines = list(loader.load_requirements([str(file_path)]))
    assert lines == ["First requirement.", "Second requirement."]

    sentences = []
    for line in lines:
        sentences.extend(loader.preprocess_text(line))
    assert sentences == []


def test_jsonl_loading_and_preprocessing(tmp_path):
    data = [
        {"text": "* The system shall reboot.", "sentence_id": "1"},
        {"text": "1. Users must change password.", "sentence_id": "2"},
    ]
    file_path = tmp_path / "reqs.jsonl"
    with open(file_path, "w", encoding="utf-8") as f:
        for obj in data:
            json.dump(obj, f)
            f.write("\n")

    loader = DataLoader()
    lines = list(loader.load_requirements([str(file_path)]))
    assert lines == data

    sentences = []
    for line in lines:
        sentences.extend(loader.preprocess_text(line["text"]))
    assert sentences == [
        "The system shall reboot.",
        "Users must change password.",
    ]


def test_jsonl_loading_with_allowed_ids(tmp_path):
    data = [
        {"text": "A", "sentence_id": "1"},
        {"text": "B", "sentence_id": "2"},
    ]
    file_path = tmp_path / "reqs.jsonl"
    with open(file_path, "w", encoding="utf-8") as f:
        for obj in data:
            json.dump(obj, f)
            f.write("\n")

    loader = DataLoader()
    lines = list(loader.load_requirements([str(file_path)], allowed_ids=["2"]))
    assert lines == [data[1]]


def test_load_requirements_warns_and_raises(tmp_path, caplog):
    loader = DataLoader()

    missing_file = tmp_path / "missing.txt"
    with caplog.at_level(logging.WARNING):
        texts = list(loader.load_requirements([str(missing_file)]))
    assert "does not exist" in caplog.text
    assert texts == []

    bad_file = tmp_path / "bad.pdf"
    bad_file.write_text("dummy")
    with pytest.raises(ValueError, match="Unsupported file extension"):
        list(loader.load_requirements([str(bad_file)]))


def test_preprocess_filters_noisy_inputs():
    loader = DataLoader()
    text = (
        "* The system shall reboot.\n"
        "1. Users must change password.\n"
        "- Bullet without verb.\n"
        "Just some text."
    )
    sentences = loader.preprocess_text(text)
    assert sentences == [
        "The system shall reboot.",
        "Users must change password.",
    ]


def test_preprocess_removes_duplicate_sentences():
    loader = DataLoader()
    text = (
        "The system shall reboot.\n"
        "Users must change password.\n"
        "The system shall reboot.\n"
        "Users must change password."
    )
    sentences = loader.preprocess_text(text)
    assert sentences == [
        "The system shall reboot.",
        "Users must change password.",
    ]


def test_preprocess_custom_keywords():
    loader = DataLoader()
    text = "Requirement: optional item."
    sentences = loader.preprocess_text(text, keywords=["requirement"])
    assert sentences == ["Requirement: optional item."]


def test_preprocess_no_keywords_argument():
    loader = DataLoader()
    text = "Logging is enabled."
    sentences = loader.preprocess_text(text, keywords=None)
    assert sentences == ["Logging is enabled."]

